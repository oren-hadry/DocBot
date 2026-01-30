"""
Word Generator - Creates professional inspection reports as Word documents
=========================================================================
Generates .docx files with structured layout, branding, and photos.
No Google account required!
"""

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openai import OpenAI

import config
from data.session_manager import ReportSession
from data.user_profile import UserProfile
from data.contacts_manager import Contact
from lang import _ as t, get_current_language

logger = logging.getLogger(__name__)

# Output directory for generated reports
REPORTS_DIR = Path(config.BASE_DIR) / "reports"
REPORTS_DIR.mkdir(exist_ok=True)


class WordGenerator:
    """Generates professional inspection reports as Word documents."""
    
    _DEFAULT_FONT = "Arial"
    
    def __init__(self):
        self.openai_client = OpenAI(api_key=config.OPENAI_API_KEY)
    
    async def generate(
        self, 
        session: ReportSession, 
        user_profile: UserProfile,
        participants: list[Contact] = None
    ) -> str:
        """
        Generate a professional inspection report as Word document.
        
        Returns: Path to the created .docx file
        """
        logger.info(f"Generating Word report for user {session.user_id}")
        
        # Step 1: Use GPT to structure the content
        structured_content = await self._structure_content(session, user_profile)
        
        # Step 2: Create Word document
        doc_path = await self._create_word_doc(
            session, 
            structured_content, 
            user_profile,
            participants or []
        )
        
        return doc_path
    
    def _resolve_report_language(self, user_profile: UserProfile, notes: str) -> str:
        """Choose a consistent report language (he/en)."""
        if user_profile.default_language and user_profile.default_language != "auto":
            return user_profile.default_language
        
        lang_code = get_current_language()
        if lang_code in ("he", "en"):
            return lang_code
        
        # Fallback heuristic based on notes
        hebrew_chars = len(re.findall(r"[\u0590-\u05FF]", notes))
        latin_chars = len(re.findall(r"[A-Za-z]", notes))
        return "he" if hebrew_chars >= latin_chars else "en"
    
    async def _structure_content(self, session: ReportSession, user_profile: UserProfile) -> dict:
        """Use GPT to organize notes into structured findings."""
        
        all_notes = session.get_all_notes()
        num_photos = len(session.photos)
        location = session.location or "Site"
        report_lang = self._resolve_report_language(user_profile, all_notes)
        report_lang_name = "Hebrew" if report_lang == "he" else "English"
        
        # If no notes, return basic structure
        if not all_notes.strip():
            return {
                "title": f"{t('doc_inspection_report')} - {location}",
                "summary": "",
                "findings": [],
                "recommendations": []
            }
        
        prompt = f"""You are a professional report writer for field inspections.

Location: {location}
Number of photos: {num_photos}
Inspector's notes:
{all_notes}

Structure this into a professional inspection report. Output JSON:
{{
    "title": "Inspection report title",
    "summary": "Brief executive summary (2-3 sentences)",
    "findings": [
        {{
            "id": 1,
            "title": "Finding title",
            "description": "Detailed description",
            "severity": "normal|important|critical"
        }}
    ],
    "recommendations": ["recommendation 1", "recommendation 2"]
}}

Rules:
- Write the report in {report_lang_name} only
- Be concise and professional
- Create one finding per distinct issue mentioned
- If severity isn't clear, use "normal"
"""
        
        try:
            response = self.openai_client.chat.completions.create(
                model=config.GPT_MODEL,
                messages=[
                    {"role": "system", "content": "You are a professional inspection report writer."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
            )
            
            structured = json.loads(response.choices[0].message.content)
            logger.info(f"Structured {len(structured.get('findings', []))} findings")
            
            return structured
            
        except Exception as e:
            logger.error(f"GPT structuring failed: {e}")
            # Return basic structure on error
            return {
                "title": f"{t('doc_inspection_report')} - {location}",
                "summary": all_notes[:200] if all_notes else "",
                "findings": [],
                "recommendations": []
            }
    
    async def _create_word_doc(
        self, 
        session: ReportSession, 
        content: dict,
        profile: UserProfile,
        participants: list[Contact]
    ) -> str:
        """Create a professional Word document report."""
        
        doc = Document()
        
        def set_run_font(run, font_name: str) -> None:
            if not run or not font_name:
                return
            run.font.name = font_name
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.get_or_add_rFonts()
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:eastAsia"), font_name)
            r_fonts.set(qn("w:cs"), font_name)
        
        def apply_paragraph_formatting(paragraph) -> None:
            if not paragraph or not paragraph.text:
                return
            is_rtl = bool(re.search(r"[\u0590-\u08FF]", paragraph.text))
            if is_rtl:
                p_pr = paragraph._p.get_or_add_pPr()
                p_pr.set(qn("w:bidi"), "1")
                if paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                set_run_font(run, self._DEFAULT_FONT)
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        location = session.location or t("doc_site_inspection")
        date_str = datetime.now().strftime("%Y-%m-%d")
        title = content.get("title", f"{t('doc_inspection_report')} - {location}")
        
        # ========== LOGO ==========
        if profile.logo_path and os.path.exists(profile.logo_path):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(profile.logo_path, width=Inches(1.5))
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")
        
        # ========== COMPANY HEADER ==========
        if profile.company_name:
            company_para = doc.add_paragraph()
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = company_para.add_run(profile.company_name)
            run.bold = True
            run.font.size = Pt(14)
        
        if profile.contact_info:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = contact_para.add_run(profile.contact_info)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Separator line
        doc.add_paragraph()
        
        # ========== REPORT TITLE ==========
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ========== DATE & LOCATION ==========
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_label = t("doc_date")
        info_para.add_run(f"{date_label}: {datetime.now().strftime('%d/%m/%Y')}")
        
        if session.location:
            loc_label = t("doc_location")
            info_para.add_run(f"   |   {loc_label}: {session.location}")
        
        doc.add_paragraph()  # Spacing
        
        # ========== PARTICIPANTS ==========
        if participants:
            part_label = t("doc_participants")
            doc.add_heading(part_label, level=2)
            
            for p in participants:
                p_text = f"• {p.name}"
                if p.organization:
                    p_text += f" ({p.organization})"
                if p.email:
                    p_text += f" - {p.email}"
                doc.add_paragraph(p_text)
            
            doc.add_paragraph()  # Spacing
        
        # ========== SUMMARY ==========
        summary = content.get("summary", "")
        if summary:
            sum_label = t("doc_summary")
            doc.add_heading(sum_label, level=2)
            doc.add_paragraph(summary)
            doc.add_paragraph()  # Spacing
        
        # ========== FINDINGS ==========
        findings = content.get("findings", [])
        if findings:
            find_label = t("doc_findings")
            doc.add_heading(find_label, level=2)
            
            for i, finding in enumerate(findings, 1):
                severity = finding.get("severity", "normal")
                severity_marker = {
                    "critical": "[!!!]",
                    "important": "[!]",
                    "normal": ""
                }.get(severity, "")
                
                f_title = finding.get("title", f"{t('doc_finding')} {i}")
                
                # Finding title
                f_para = doc.add_paragraph()
                run = f_para.add_run(f"{i}. {f_title} {severity_marker}")
                run.bold = True
                
                # Severity color
                if severity == "critical":
                    run.font.color.rgb = RGBColor(180, 0, 0)
                elif severity == "important":
                    run.font.color.rgb = RGBColor(200, 150, 0)
                
                # Finding description
                f_desc = finding.get("description", "")
                if f_desc:
                    desc_para = doc.add_paragraph(f_desc)
                    desc_para.paragraph_format.left_indent = Inches(0.3)
                
                doc.add_paragraph()  # Spacing between findings
        
        # ========== RECOMMENDATIONS ==========
        recommendations = content.get("recommendations", [])
        if recommendations:
            rec_label = t("doc_recommendations")
            doc.add_heading(rec_label, level=2)
            
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}")
            
            doc.add_paragraph()  # Spacing
        
        # ========== PHOTOS ==========
        if session.photos:
            photos_label = t("doc_photos")
            doc.add_heading(photos_label, level=2)
            
            photo_label = t("doc_photo")
            for i, photo_path in enumerate(session.photos, 1):
                if os.path.exists(photo_path):
                    try:
                        # Photo label
                        doc.add_paragraph(f"{photo_label} {i}:")
                        
                        # Insert photo
                        photo_para = doc.add_paragraph()
                        photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = photo_para.add_run()
                        run.add_picture(photo_path, width=Inches(5))
                        
                        doc.add_paragraph()  # Spacing
                        
                    except Exception as e:
                        logger.warning(f"Could not add photo {i}: {e}")
                        doc.add_paragraph(f"[{photo_label} {i} - {t('doc_photo_error')}]")
        
        # ========== FOOTER WITH PAGE NUMBERS ==========
        # Note: python-docx has limited footer support, adding simple footer text
        footer_section = doc.sections[0]
        footer = footer_section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f"{profile.get_company_display()} | {t('doc_generated_by_docbot')}")
        
        # Normalize fonts and RTL across all paragraphs
        for paragraph in doc.paragraphs:
            apply_paragraph_formatting(paragraph)
        for section in doc.sections:
            for paragraph in section.footer.paragraphs:
                apply_paragraph_formatting(paragraph)
        
        # ========== SAVE DOCUMENT ==========
        # Create unique filename
        safe_location = "".join(c for c in location if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_location = safe_location[:30]  # Limit length
        
        filename = f"Report_{safe_location}_{date_str}_{session.user_id}.docx"
        filepath = REPORTS_DIR / filename
        
        doc.save(str(filepath))
        logger.info(f"Word report saved: {filepath}")
        
        return str(filepath)


# Singleton instance
word_generator = WordGenerator()
