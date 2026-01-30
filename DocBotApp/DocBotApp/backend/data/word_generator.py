import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

from data.contacts import contacts_manager
from data.report_manager import ReportSession
from data.storage import user_reports_dir


def generate_report_docx(session: ReportSession) -> str:
    doc = Document()

    doc.add_heading(session.title, level=1)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    if session.location:
        doc.add_paragraph(f"Location: {session.location}")

    if session.attendees:
        attendees = contacts_manager.get_contacts_by_ids(session.user_id, session.attendees)
        doc.add_paragraph("Attendees:")
        for c in attendees:
            line = f"- {c.name}"
            if c.email:
                line += f" ({c.email})"
            doc.add_paragraph(line)

    if session.distribution_list:
        recipients = contacts_manager.get_contacts_by_ids(session.user_id, session.distribution_list)
        doc.add_paragraph("Distribution:")
        for c in recipients:
            line = f"- {c.name}"
            if c.email:
                line += f" ({c.email})"
            doc.add_paragraph(line)

    doc.add_heading("Findings", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Notes"
    hdr_cells[3].text = "Photo"

    for item in session.items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.number
        row_cells[1].text = item.description
        row_cells[2].text = item.notes or ""
        photos = [p for p in session.photos if p.item_id == item.id]
        if photos:
            try:
                row_cells[3].paragraphs[0].add_run().add_picture(
                    photos[0].file_path, width=Inches(1.5)
                )
            except Exception:
                row_cells[3].text = "Photo attached"

    if session.photos:
        doc.add_page_break()
        doc.add_heading("Photo Appendix", level=2)
        for photo in session.photos:
            item_number = "-"
            for item in session.items:
                if item.id == photo.item_id:
                    item_number = item.number
                    break
            doc.add_paragraph(f"Item {item_number}")
            try:
                doc.add_picture(photo.file_path, width=Inches(5))
            except Exception:
                doc.add_paragraph("Photo unavailable")

    reports_dir = user_reports_dir(session.user_id)
    filename = f"Report_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.docx"
    filepath = reports_dir / filename
    doc.save(str(filepath))
    return str(filepath)
