import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from data.contacts import contacts_manager
from data.report_manager import ReportSession
from data.storage import user_reports_dir
from data.user_auth import user_auth

HEBREW_TITLES = {
    "INSPECTION_REPORT": "דוח פיקוח",
    "VISIT_SUMMARY": "סיכום ביקור",
    "HOME_ORGANIZER_REPORT": "דוח סידור בית",
    "QUOTE": "הצעת מחיר",
}


def _contains_hebrew(text: str) -> bool:
    return any("\u0590" <= ch <= "\u05FF" for ch in text)


def _set_run_font(run, font_name: str, is_rtl: bool = False) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    if is_rtl:
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        rpr.append(rtl)


def _set_paragraph_rtl(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    ppr.append(bidi)
    paragraph.alignment = 2  # right


def _format_paragraph(paragraph, font_name: str, force_rtl: bool = False) -> None:
    text = paragraph.text or ""
    is_rtl = force_rtl or _contains_hebrew(text)
    if is_rtl:
        _set_paragraph_rtl(paragraph)
    for run in paragraph.runs:
        _set_run_font(run, font_name, is_rtl=is_rtl)


def _set_table_rtl(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tbl_pr.append(bidi)


def _format_table(table, font_name: str, force_rtl: bool = False) -> None:
    if force_rtl:
        _set_table_rtl(table)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _format_paragraph(paragraph, font_name, force_rtl=force_rtl)


def generate_report_docx(session: ReportSession, logo_path: Optional[str] = None) -> str:
    doc = Document()
    
    # Set document-wide RTL properties
    for section in doc.sections:
        if any(_contains_hebrew(v or "") for v in [session.title, session.location]):
            sect_pr = section._sectPr
            bidi_sect = OxmlElement("w:bidi")
            bidi_sect.set(qn("w:val"), "1")
            sect_pr.append(bidi_sect)

    font_name = "Arial"
    doc.styles["Normal"].font.name = font_name

    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(1.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = 1 # center
        except Exception:
            pass

    has_hebrew = any(
        _contains_hebrew(value or "")
        for value in [
            session.title,
            session.location,
            *[i.description for i in session.items],
            *[i.notes for i in session.items],
        ]
    )

    title_text = HEBREW_TITLES.get(session.template_key, session.title) if has_hebrew else session.title
    labels = {
        "date": "תאריך" if has_hebrew else "Date",
        "location": "מיקום" if has_hebrew else "Location",
        "attendees": "נוכחים" if has_hebrew else "Attendees",
        "distribution": "תפוצה" if has_hebrew else "Distribution",
        "findings": "ממצאים" if has_hebrew else "Findings",
        "description": "תיאור" if has_hebrew else "Description",
        "notes": "הערות" if has_hebrew else "Notes",
        "photo": "תמונה" if has_hebrew else "Photo",
        "item": "סעיף" if has_hebrew else "Item",
    }

    doc.add_heading(title_text, level=1)

    # Date section
    doc.add_heading(labels['date'], level=2)
    doc.add_paragraph(datetime.now().strftime('%d/%m/%Y'))

    if session.location:
        doc.add_heading(labels['location'], level=2)
        doc.add_paragraph(session.location)

    if session.attendees:
        attendees = contacts_manager.get_contacts_by_ids(session.user_id, session.attendees)
        doc.add_heading(labels['attendees'], level=2)
        for c in attendees:
            line = f"- {c.name}"
            if c.email:
                line += f" ({c.email})"
            doc.add_paragraph(line)

    if session.distribution_list:
        recipients = contacts_manager.get_contacts_by_ids(session.user_id, session.distribution_list)
        doc.add_heading(labels['distribution'], level=2)
        for c in recipients:
            line = f"- {c.name}"
            if c.email:
                line += f" ({c.email})"
            doc.add_paragraph(line)

    doc.add_heading(labels["findings"], level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"  # Add visible borders
    table.alignment = 2 # Right align the table on the page
    table.autofit = False
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(3.0)
    table.columns[3].width = Inches(1.5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = labels["description"]
    hdr_cells[2].text = labels["notes"]
    hdr_cells[3].text = labels["photo"]
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for item in session.items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.number
        row_cells[1].text = item.description
        row_cells[2].text = item.notes or ""
        photos = [p for p in session.photos if p.item_id == item.id]
        if photos:
            try:
                row_cells[3].paragraphs[0].add_run().add_picture(
                    photos[0].file_path, width=Inches(1.5)
                )
            except Exception:
                row_cells[3].text = labels["photo"]

    # Signature section
    user = user_auth.get_by_id(session.user_id)
    if user and user.full_name:
        doc.add_paragraph() # Spacer
        
        if user.signature_path and os.path.exists(user.signature_path):
            try:
                doc.add_picture(user.signature_path, width=Inches(1.2))
            except Exception:
                pass
        
        doc.add_paragraph("-" * 20)
        
        written_by = "מסמך זה נכתב על ידי" if has_hebrew else "This document was written by"
        p = doc.add_paragraph()
        run = p.add_run(f"{written_by} {user.full_name}")
        run.bold = True
        
        if user.role_title:
            doc.add_paragraph(user.role_title)
        
        contact_info = []
        if user.company_name:
            contact_info.append(user.company_name)
        if user.phone_contact:
            contact_info.append(user.phone_contact)
        elif user.phone:
            contact_info.append(user.phone)
            
        if contact_info:
            doc.add_paragraph(" | ".join(contact_info))

    for paragraph in doc.paragraphs:
        _format_paragraph(paragraph, font_name, force_rtl=has_hebrew)
    for table in doc.tables:
        _format_table(table, font_name, force_rtl=has_hebrew)

    reports_dir = user_reports_dir(session.user_id)
    filename = f"Report_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.docx"
    filepath = reports_dir / filename
    doc.save(str(filepath))
    return str(filepath)
